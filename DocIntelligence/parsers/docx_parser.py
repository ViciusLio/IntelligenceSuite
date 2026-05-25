"""Parser DOCX — heading-aware + estrazione tabelle in markdown."""

from __future__ import annotations
import re
from pathlib import Path

from intelligence_core.chunk import make_chunk

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def can_parse(path: Path) -> bool:
    return path.suffix.lower() in {".docx", ".doc"}


def parse_file(path: Path, root: Path) -> list[dict]:
    """Ritorna chunk da DOCX. Fallback a chunk raw."""
    rel = str(path.relative_to(root)).replace("\\", "/")
    stem = path.stem

    if not HAS_DOCX:
        return [make_chunk(
            domain="doc", type_="file", locator=stem,
            text=f"File: {path.name} (in {rel})\n---\n[python-docx non installato]",
            source=rel, language="docx",
            metadata={"page_start": 1, "page_end": 1, "heading_path": [stem]},
        )]

    try:
        return _parse_docx(path, rel, stem)
    except Exception as e:
        return [make_chunk(
            domain="doc", type_="file", locator=stem,
            text=f"File: {path.name} (in {rel})\n---\n[Errore parsing: {e}]",
            source=rel, language="docx",
            metadata={"page_start": 1, "page_end": 1, "heading_path": [stem]},
        )]


def _parse_docx(path: Path, rel: str, stem: str) -> list[dict]:
    doc = Document(str(path))
    chunks = []
    heading_path: list[str] = []
    body_lines: list[str] = []
    table_counter = 0

    def flush_section():
        body = "\n".join(body_lines).strip()
        heading = heading_path[-1] if heading_path else stem
        # Crea comunque un chunk per sezioni senza body (es. heading seguito da tabella)
        if len(body) < 30:
            if not heading_path:
                return
            body = f"[Sezione senza corpo]"
        text = (
            f"Section: {heading} (in {path.name})\n"
            f"Path: {' > '.join(heading_path) if heading_path else stem}\n"
            f"---\n{body}"
        )
        locator = f"{stem}.{_sanitize(' > '.join(heading_path))}"
        chunks.append(make_chunk(
            domain="doc", type_="section", locator=locator,
            text=text[:7500], source=rel, language="docx",
            metadata={
                "page_start":   1,
                "page_end":     1,
                "heading_path": list(heading_path),
                "word_count":   len(body.split()),
            },
        ))

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragrafo
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            style_name = para.style.name if para.style else ""
            text_content = para.text.strip()
            if not text_content:
                continue

            if "Heading 1" in style_name or style_name.startswith("Title"):
                flush_section()
                body_lines.clear()
                heading_path = [text_content]
            elif "Heading 2" in style_name:
                flush_section()
                body_lines.clear()
                heading_path = heading_path[:1] + [text_content]
            elif "Heading 3" in style_name:
                flush_section()
                body_lines.clear()
                heading_path = heading_path[:2] + [text_content]
            else:
                body_lines.append(text_content)

        elif tag == "tbl":
            flush_section()
            body_lines.clear()

            tbl_chunk = _extract_table(element, doc, path.name, rel, stem, table_counter)
            if tbl_chunk:
                chunks.append(tbl_chunk)
            table_counter += 1

    flush_section()

    if not chunks:
        # Fallback: testo completo
        full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:3000]
        chunks.append(make_chunk(
            domain="doc", type_="file", locator=stem,
            text=f"File: {path.name} (in {rel})\n---\n{full}",
            source=rel, language="docx",
            metadata={"page_start": 1, "page_end": 1, "heading_path": [stem]},
        ))

    return chunks


def _extract_table(element, doc, filename: str, rel: str, stem: str, idx: int) -> dict | None:
    """Converte una tabella DOCX in chunk markdown."""
    try:
        tbl = None
        for t in doc.tables:
            if t._element is element:
                tbl = t
                break
        if tbl is None:
            return None

        rows = tbl.rows
        if not rows:
            return None

        header = [cell.text.strip() for cell in rows[0].cells]
        separator = ["---"] * len(header)
        data_rows = [[cell.text.strip() for cell in row.cells] for row in rows[1:]]

        def row_to_md(cells):
            return "| " + " | ".join(cells) + " |"

        md_lines = [
            row_to_md(header),
            row_to_md(separator),
        ]
        for row in data_rows:
            md_lines.append(row_to_md(row))

        md_table = "\n".join(md_lines)
        text = (
            f"Table (in {filename})\n"
            f"Columns: {', '.join(header)}\n"
            f"---\n{md_table}"
        )

        return make_chunk(
            domain="doc", type_="table", locator=f"{stem}.table_{idx}",
            text=text[:7500], source=rel, language="docx",
            metadata={
                "page_start": 1, "page_end": 1,
                "heading_path": [stem, f"Table {idx}"],
                "columns": header,
                "row_count": len(data_rows),
            },
        )
    except Exception:
        return None


def _sanitize(s: str) -> str:
    return re.sub(r"[^\w]", "_", s.lower())[:60]
