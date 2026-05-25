"""Genera i documenti di esempio in docs/ per i test dei parser."""

from fpdf import FPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill
from pathlib import Path

OUT = Path("docs")
OUT.mkdir(exist_ok=True)


def write_line(pdf, line):
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 5, line.encode("latin-1", "replace").decode("latin-1"))


def make_pdf(filename, title, rgb, sections):
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_fill_color(*rgb)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 14, title, new_x="LMARGIN", new_y="NEXT", align="C", fill=True)
    pdf.ln(4)
    pdf.set_text_color(0, 0, 0)
    for sec_title, lines in sections:
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(235, 235, 235)
        pdf.set_x(pdf.l_margin)
        pdf.cell(0, 7, sec_title, new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10)
        for line in lines:
            write_line(pdf, line)
        pdf.ln(3)
    pdf.output(str(OUT / filename))
    print(f"  {filename}")


make_pdf(
    "api_reference_v2.pdf",
    "API Reference - IntelligenceSuite v2.1",
    (30, 60, 114),
    [
        ("1. Panoramica", [
            "IntelligenceSuite espone una REST API su porta 8080.",
            "Base URL: http://localhost:8080/api/v1",
            "Autenticazione via JWT Bearer Token.",
        ]),
        ("2. Autenticazione", [
            "POST /api/v1/auth/token con client_id e client_secret.",
            "Response: access_token, expires_in 3600 secondi.",
            "Ruoli: admin, reader, indexer.",
            "Bearer token obbligatorio in header Authorization: Bearer TOKEN.",
        ]),
        ("3. Endpoint /query", [
            "POST /api/v1/query",
            "Parametri: question (obbligatorio), domains, top_k, min_score.",
            "Response: answer, sources, confidence, escalated, latency_ms.",
            "Errori: 400 query vuota, 401 token scaduto, 429 rate limit.",
        ]),
        ("4. Endpoint /ingest", [
            "POST /api/v1/ingest",
            "Parametri: path, domain (code|doc|auto), incremental.",
            "Response: job_id, status, files_queued, estimated_seconds.",
        ]),
        ("5. Rate Limiting e SLA", [
            "Free: 100 query/giorno, 10 query/minuto.",
            "Team: 1000 query/giorno, 100 query/minuto.",
            "Enterprise: illimitato, 500/minuto.",
            "SLA Enterprise: P50 < 300ms, P99 < 2000ms, uptime 99.9%, Hit@5 > 85%.",
        ]),
    ],
)

make_pdf(
    "deploy_procedure_v3.pdf",
    "Procedura Operativa: Deploy in Produzione",
    (20, 100, 60),
    [
        ("1. Scopo", [
            "Deploy di IntelligenceSuite in EU-WEST-1.",
            "Approvazione richiesta: Tech Lead e Product Owner.",
            "Versione 3.2 - Team Platform - 2026-04-15",
        ]),
        ("2. Pre-Requisiti", [
            "[ ] CI verde su GitHub Actions",
            "[ ] Code review da 2 reviewer",
            "[ ] Changelog aggiornato",
            "[ ] Backup vector store completato",
            "[ ] Monitoring attivo su Grafana",
            "[ ] Nessun incident P1/P2 su PagerDuty",
        ]),
        ("3. Procedura", [
            "STEP 1: Maintenance mode ON.",
            "STEP 2: Attendere esaurimento connessioni.",
            "STEP 3: kubectl set image deployment/intelligence-suite app=IMAGE:VERSION",
            "STEP 4: Smoke test automatico 60 secondi.",
            "STEP 5: Maintenance mode OFF.",
            "STEP 6: Monitorare Grafana 15 minuti. Allarme se error rate > 1%.",
        ]),
        ("4. Rollback", [
            "Entro 30 minuti: kubectl rollout undo deployment/intelligence-suite -n production",
            "Oltre 30 minuti: kubectl set image con versione precedente.",
            "Post-rollback: incident PagerDuty, notifica #incidents, post-mortem 48 ore.",
        ]),
        ("5. Contatti", [
            "Platform on-call: +39 02 XXXX XXXX",
            "Tech Lead: mario.rossi@company.com",
            "CTO: cto@company.com",
        ]),
    ],
)

doc = Document()
t = doc.add_heading("ADR-007: Scelta Vector Store per Intelligence Suite", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading("Contesto", 1)
doc.add_paragraph(
    "Intelligence Suite richiede un vector store on-premise per GDPR compliance, "
    "scalabile da 100K a 10M chunk, con costi operativi minimi."
)
doc.add_heading("Decisione", 1)
doc.add_paragraph(
    "Adottiamo pgvector (PostgreSQL) come default per produzione, "
    "ChromaDB per sviluppo/single-node. "
    "La scelta e' stata adottata per motivi di compliance GDPR e costi operativi."
)
doc.add_heading("Opzioni Valutate", 1)
tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Table Grid"
for i, h in enumerate(["Soluzione", "Self-hosted", "Scala 10M", "Costo ops", "Verdict"]):
    tbl.rows[0].cells[i].text = h
    tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for rd in [
    ["pgvector",  "SI",  "SI",      "Basso (PostgreSQL)", "SCELTO"],
    ["ChromaDB",  "SI",  "Parziale","Minimo",              "OK dev"],
    ["Weaviate",  "SI",  "SI",      "Alto",                "Overkill"],
    ["Pinecone",  "NO",  "SI",      "Pay-per-use",         "Non compliance"],
    ["Qdrant",    "SI",  "SI",      "Medio",               "Candidato v2"],
]:
    row = tbl.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
doc.add_heading("Conseguenze Positive", 1)
for p in [
    "Zero infrastruttura nuova per chi ha PostgreSQL.",
    "pgvector HNSW: ricerca ANN su 1M vettori in < 10ms.",
    "SQL standard per query ibride.",
]:
    doc.add_paragraph(p, style="List Bullet")
doc.add_heading("Conseguenze Negative", 1)
for p in [
    "Richiede PostgreSQL >= 14 con estensione pgvector installata.",
    "Setup iniziale piu' complesso rispetto a ChromaDB.",
]:
    doc.add_paragraph(p, style="List Bullet")
doc.add_heading("Metriche di Successo", 1)
for p in [
    "Latenza P99 < 50ms su 500K chunk.",
    "Zero incidenti vector store in produzione.",
]:
    doc.add_paragraph(p, style="List Bullet")
doc.save(str(OUT / "ADR-007-vector-store.docx"))
print("  ADR-007-vector-store.docx")

wb = openpyxl.Workbook()
hf = PatternFill("solid", fgColor="1E3C72")
hfont = Font(color="FFFFFF", bold=True, size=11)
gf = PatternFill("solid", fgColor="C6EFCE")
tf = PatternFill("solid", fgColor="D9E1F2")
ws1 = wb.active
ws1.title = "KPI Dashboard"
ws1["A1"] = "Intelligence Suite - KPI Dashboard"
ws1["A1"].font = Font(bold=True, size=14)
ws1.append([])
hdrs = ["KPI", "Target", "Attuale", "Trend", "Stato", "Note"]
ws1.append(hdrs)
for col, h in enumerate(hdrs, 1):
    c = ws1.cell(3, col)
    c.fill = hf
    c.font = hfont
for kpi in [
    ["Hit@1 (CodeIntelligence)",  "> 60%",  "67%",  "up +3%",      "OK",   "150 query reali"],
    ["Hit@5 (CodeIntelligence)",  "> 85%",  "88%",  "stabile",     "OK",   "Soglia minima"],
    ["MRR (CodeIntelligence)",    "> 0.70", "0.74", "up +0.02",    "OK",   "Mean Reciprocal Rank"],
    ["NDCG@5 (CodeIntelligence)", "> 0.75", "0.79", "up +0.01",    "OK",   ""],
    ["Hit@1 (DocIntelligence)",   "> 55%",  "N/A",  "-",           "TODO", "Parser in sviluppo"],
    ["Hit@5 (DocIntelligence)",   "> 80%",  "N/A",  "-",           "TODO", ""],
    ["Latenza P50 locale (ms)",   "< 300",  "187",  "down -20",    "OK",   "RTX 4090"],
    ["Latenza P99 locale (ms)",   "< 1000", "780",  "stabile",     "OK",   ""],
    ["Tasso escalation",          "< 15%",  "11%",  "down -2%",    "OK",   ""],
    ["Tasso allucinazione",       "< 5%",   "3.2%", "down -0.8%",  "OK",   "Campione manuale"],
]:
    ws1.append(kpi)
    r = ws1.max_row
    for col in range(1, 7):
        ws1.cell(r, col).fill = tf if kpi[4] == "TODO" else gf
for col, w in zip("ABCDEF", [35, 14, 12, 14, 10, 40]):
    ws1.column_dimensions[col].width = w

ws2 = wb.create_sheet("Benchmark Queries")
h2 = ["ID", "Dominio", "Query", "Risposta Attesa", "Chunk Target", "Difficolta", "Categoria"]
ws2.append(h2)
for col, h in enumerate(h2, 1):
    c = ws2.cell(1, col)
    c.fill = hf
    c.font = hfont
for q in [
    ["Q001","code","Dove viene gestita l autenticazione JWT?","validate_token in app/auth.py","code::function::auth.validate_token","facile","retrieval"],
    ["Q002","code","Quali endpoint richiedono ruolo admin?","@require_role admin","code::function::routes.admin_endpoints","media","reasoning"],
    ["Q009","doc","Come ottenere un token API?","POST /api/v1/auth/token","doc::section::APIRef.Autenticazione","facile","retrieval"],
    ["Q010","doc","Rate limit piano Team?","1000/giorno 100/minuto","doc::section::APIRef.RateLimiting","facile","retrieval"],
    ["Q011","doc","Cosa fare se deploy fallisce entro 30 minuti?","kubectl rollout undo","doc::section::DeployProc.Rollback","media","reasoning"],
    ["Q012","doc","Perche pgvector invece di Pinecone?","Non compliance GDPR","doc::section::ADR007.Decisione","difficile","reasoning"],
]:
    ws2.append(q)

wb.save(str(OUT / "intelligence_suite_kpi_benchmark.xlsx"))
print("  intelligence_suite_kpi_benchmark.xlsx")
print("Tutti i documenti generati.")
